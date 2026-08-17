"""
BiTremplin - API de traduction (version relais vers Kaggle+ngrok, + traduction
de documents Word/PDF/txt, avec streaming SSE)
====================================================================
L'Inference API Serverless de Hugging Face ne supporte pas les modeles
fine-tunes personnels non adoptes par un fournisseur (erreur StopIteration
constatee en test). Cette version relaie donc les requetes de traduction
vers une session Kaggle qui fait tourner le vrai modele (voir
run_api_kaggle_ngrok.py), exposee via un tunnel ngrok.

Render reste le point d'entree stable pour le frontend (URL fixe), mais le
calcul reel depend d'une session Kaggle active -> renseigne KAGGLE_API_URL
(variable d'environnement sur Render) a chaque redemarrage de la session
Kaggle, avec la nouvelle URL ngrok affichee par le script.
"""

import asyncio
import io
import json
import os
import re
from pathlib import Path
from typing import Optional, List

import httpx
from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

MODEL_NAME = os.environ.get("MODEL_NAME", "Expendadeur/nllb-kin-fr")

# URL ngrok de la session Kaggle active (voir run_api_kaggle_ngrok.py).
# A METTRE A JOUR sur Render (Settings > Environment) a chaque redemarrage
# de la session Kaggle, car l'URL ngrok change a chaque fois (sauf domaine
# fixe reserve). Sans ca, /translate repond une erreur claire plutot que de
# planter silencieusement.
KAGGLE_API_URL = os.environ.get("KAGGLE_API_URL", "")

ADMIN_TOKEN = os.environ.get("ADMIN_TOKEN", "change-me")

BASE_DIR = Path(__file__).parent
LANG_FILE = BASE_DIR / "languages.json"
ADMIN_FILE = BASE_DIR / "admins.json"

REQUEST_TIMEOUT_S = 60  # le modele + generation peut prendre du temps sur CPU Kaggle

# Taille max d'un morceau de texte envoye en une seule traduction. Le modele a
# ete entraine avec max_length=128 tokens ; on reste prudent en caracteres
# pour eviter les troncatures silencieuses qui degraderaient la qualite.
CHUNK_MAX_CHARS = 400
MAX_DOCUMENT_CHARS = 200_000  # garde-fou anti-abus (~40-60 pages) sur un service gratuit

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

app = FastAPI(title="BiTremplin Translation API (relais Kaggle)", version="1.2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if not KAGGLE_API_URL:
    print("[BiTremplin] ATTENTION : KAGGLE_API_URL non definie. /translate renverra une erreur claire tant qu'elle ne l'est pas.")
else:
    print(f"[BiTremplin] Relais configure vers : {KAGGLE_API_URL}")


# ---------------------------------------------------------------------------
# "Base de donnees" fichier JSON (langues + admins)
# ---------------------------------------------------------------------------

def _read_json(path: Path, default):
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return default

def _write_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def load_languages() -> List[dict]:
    return _read_json(LANG_FILE, [])

def save_languages(langs: List[dict]):
    _write_json(LANG_FILE, langs)

def load_admins() -> dict:
    return _read_json(ADMIN_FILE, {"emails": []})

def save_admins(data: dict):
    _write_json(ADMIN_FILE, data)


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class TranslateRequest(BaseModel):
    text: str
    source: str
    target: str

class TranslateResponse(BaseModel):
    translation: str
    source: str
    target: str

class LanguageItem(BaseModel):
    code: str
    label: str
    flag: Optional[str] = None
    ttsLocale: Optional[str] = None
    modelCode: Optional[str] = None

class AdminEmailItem(BaseModel):
    email: str

class BuildDocumentRequest(BaseModel):
    paragraphs: List[str]
    output_format: str = "docx"  # "docx" ou "txt"


def check_admin(x_admin_token: str = Header(default="")):
    if not ADMIN_TOKEN or x_admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=403, detail="Acces admin refuse")
    return True


def check_languages_allowed(source: str, target: str) -> dict:
    """Verifie que les deux langues sont autorisees et renvoie le mapping
    complet code -> config (reutilise par /translate et /translate-document)."""
    langs_by_code = {lang["code"]: lang for lang in load_languages()}
    if source not in langs_by_code:
        raise HTTPException(status_code=403, detail=f"Langue source non autorisee : {source}")
    if target not in langs_by_code:
        raise HTTPException(status_code=403, detail=f"Langue cible non autorisee : {target}")
    return langs_by_code


def translate_text(text: str, source: str, target: str) -> str:
    """Traduction d'un morceau de texte unique via le relais Kaggle. Utilisee
    a la fois par /translate et par /translate-document-stream (chunk par
    chunk)."""
    if not KAGGLE_API_URL:
        raise HTTPException(
            status_code=503,
            detail="KAGGLE_API_URL non configuree sur Render. Lance run_api_kaggle_ngrok.py "
                   "dans une session Kaggle, puis colle l'URL ngrok affichee dans les "
                   "variables d'environnement Render (Settings > Environment).",
        )
    try:
        resp = httpx.post(
            f"{KAGGLE_API_URL.rstrip('/')}/translate",
            json={"text": text, "source": source, "target": target},
            timeout=REQUEST_TIMEOUT_S,
        )
        resp.raise_for_status()
        return resp.json()["translation"]
    except httpx.ConnectError:
        raise HTTPException(
            status_code=503,
            detail="Impossible de joindre la session Kaggle. Verifie qu'elle est bien "
                   "active et que KAGGLE_API_URL est a jour (l'URL ngrok change a "
                   "chaque redemarrage de la session).",
        )
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="La session Kaggle n'a pas repondu a temps.")
    except httpx.HTTPStatusError as e:
        raise HTTPException(
            status_code=502,
            detail=f"Erreur renvoyee par la session Kaggle ({e.response.status_code}) : {e.response.text}",
        )


# ---------------------------------------------------------------------------
# Extraction de texte depuis un document (txt / docx / pdf)
# ---------------------------------------------------------------------------

def extract_paragraphs(filename: str, raw_bytes: bytes) -> List[str]:
    """Renvoie une liste de paragraphes (texte brut) extraits du document,
    quel que soit son format. Chaque entree devient un paragraphe distinct
    dans le document traduit en sortie."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "txt":
        text = raw_bytes.decode("utf-8", errors="replace")
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    elif ext == "docx":
        try:
            import docx
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx non installe cote serveur.")
        document = docx.Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    elif ext == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise HTTPException(status_code=500, detail="pypdf non installe cote serveur.")
        reader = PdfReader(io.BytesIO(raw_bytes))
        paragraphs = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            paragraphs.extend(p.strip() for p in page_text.split("\n") if p.strip())

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Format non supporte : .{ext}. Formats acceptes : .txt, .docx, .pdf",
        )

    if not paragraphs:
        raise HTTPException(status_code=400, detail="Aucun texte extrait du document (fichier vide ou illisible).")

    total_chars = sum(len(p) for p in paragraphs)
    if total_chars > MAX_DOCUMENT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"Document trop volumineux ({total_chars} caracteres, max {MAX_DOCUMENT_CHARS}).",
        )

    return paragraphs


def split_into_chunks(paragraph: str, max_chars: int = CHUNK_MAX_CHARS) -> List[str]:
    """Decoupe un paragraphe en morceaux traduisibles individuellement, en
    coupant sur les frontieres de phrases plutot qu'au milieu d'un mot."""
    if len(paragraph) <= max_chars:
        return [paragraph]

    sentences = re.split(r"(?<=[.!?])\s+", paragraph)
    chunks, current = [], ""
    for sentence in sentences:
        if len(current) + len(sentence) + 1 <= max_chars:
            current = f"{current} {sentence}".strip()
        else:
            if current:
                chunks.append(current)
            # Phrase elle-meme trop longue -> coupe brutalement en dernier recours
            current = sentence if len(sentence) <= max_chars else sentence[:max_chars]
    if current:
        chunks.append(current)
    return chunks


def translate_paragraphs(paragraphs: List[str], source: str, target: str) -> List[str]:
    translated = []
    for paragraph in paragraphs:
        chunks = split_into_chunks(paragraph)
        translated_chunks = [translate_text(chunk, source, target) for chunk in chunks]
        translated.append(" ".join(translated_chunks))
    return translated


def build_output_file(paragraphs: List[str], output_format: str):
    """Construit le fichier de sortie a partir des paragraphes traduits.
    Renvoie (contenu binaire, nom de fichier, media type)."""
    if output_format == "docx":
        try:
            import docx
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx non installe cote serveur.")
        document = docx.Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        buffer = io.BytesIO()
        document.save(buffer)
        return (
            buffer.getvalue(),
            "traduction.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:  # txt par defaut - le plus simple et le plus fiable dans tous les cas
        content = "\n\n".join(paragraphs).encode("utf-8")
        return content, "traduction.txt", "text/plain; charset=utf-8"


def sse_event(event: str, data: dict) -> str:
    """Formatte une ligne au format SSE (text/event-stream)."""
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ---------------------------------------------------------------------------
# Routes publiques
# ---------------------------------------------------------------------------

@app.get("/")
def root():
    return {
        "status": "ok",
        "model": MODEL_NAME,
        "mode": "relais-kaggle",
        "kaggle_configured": bool(KAGGLE_API_URL),
    }


@app.get("/languages", response_model=List[LanguageItem])
def get_languages():
    return load_languages()


@app.get("/is-admin")
def is_admin(email: str):
    data = load_admins()
    return {"is_admin": email.lower() in [e.lower() for e in data["emails"]]}


@app.post("/translate", response_model=TranslateResponse)
def translate(req: TranslateRequest):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Texte vide")
    if len(req.text) > 5000:
        raise HTTPException(status_code=400, detail="Texte trop long (max 5000 caracteres)")

    check_languages_allowed(req.source, req.target)
    translation = translate_text(req.text, req.source, req.target)
    return TranslateResponse(translation=translation, source=req.source, target=req.target)


@app.post("/translate-document")
async def translate_document(
    file: UploadFile = File(...),
    source: str = Form(...),
    target: str = Form(...),
    output_format: str = Form("docx"),  # "docx" ou "txt"
):
    """Version NON streamee (conservee pour compatibilite / usage simple sans
    affichage progressif) : traduit le document entier et renvoie directement
    le fichier final. Pour l'affichage streaming cote frontend, preferer
    /translate-document-stream + /build-document (voir plus bas).

    NOTE : le PDF ne conserve pas sa mise en page d'origine (tableaux,
    images, colonnes) - seul le texte est extrait et retraduit dans un
    document simple. Pour un .docx en entree, la structure en paragraphes
    est preservee ; polices, styles et images ne le sont pas.
    """
    if output_format not in ("docx", "txt"):
        raise HTTPException(status_code=400, detail="output_format doit etre 'docx' ou 'txt'")

    check_languages_allowed(source, target)

    raw_bytes = await file.read()
    paragraphs = extract_paragraphs(file.filename or "document.txt", raw_bytes)
    translated_paragraphs = translate_paragraphs(paragraphs, source, target)
    content, out_filename, media_type = build_output_file(translated_paragraphs, output_format)

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
    )


@app.post("/translate-document-stream")
async def translate_document_stream(
    file: UploadFile = File(...),
    source: str = Form(...),
    target: str = Form(...),
):
    """
    Traduction de document en streaming SSE (text/event-stream) :

      event: source     -> {"paragraphs": [...]}               (texte original, tout de suite)
      event: paragraph  -> {"index": i, "translation": "..."}  (un par paragraphe traduit)
      event: error      -> {"index": i, "detail": "..."}       (si un paragraphe echoue)
      event: done        -> {"count": n}                        (fin du flux)

    Le frontend affiche `source` immediatement a gauche, puis ajoute chaque
    `paragraph` a droite au fur et a mesure (effet machine a ecrire). Une
    fois le flux termine, il dispose de tous les paragraphes traduits et
    peut appeler /build-document pour generer le fichier telechargeable,
    sans re-traduire.
    """
    check_languages_allowed(source, target)

    raw_bytes = await file.read()
    paragraphs = extract_paragraphs(file.filename or "document.txt", raw_bytes)

    async def event_generator():
        # 1) Texte source complet, envoye immediatement
        yield sse_event("source", {"paragraphs": paragraphs})

        # 2) Traduction paragraphe par paragraphe, streamee au fur et a mesure
        for i, paragraph in enumerate(paragraphs):
            try:
                chunks = split_into_chunks(paragraph)
                # translate_text() est bloquant (httpx sync) -> on le passe
                # dans un thread pour ne pas geler l'event loop et laisser
                # les octets deja generes partir vers le client.
                translated_chunks = [
                    await asyncio.to_thread(translate_text, chunk, source, target)
                    for chunk in chunks
                ]
                translation = " ".join(translated_chunks)
                yield sse_event("paragraph", {"index": i, "translation": translation})
            except HTTPException as e:
                yield sse_event("error", {"index": i, "detail": e.detail})
                # On continue avec les paragraphes suivants plutot que de
                # couper tout le flux pour un seul paragraphe en echec.

        yield sse_event("done", {"count": len(paragraphs)})

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # evite le buffering cote proxy (Render)
        },
    )


@app.post("/build-document")
def build_document(req: BuildDocumentRequest):
    """
    Construit le fichier final a partir de paragraphes DEJA traduits (recus
    du flux /translate-document-stream cote client). Ne retraduit rien.
    """
    if req.output_format not in ("docx", "txt"):
        raise HTTPException(status_code=400, detail="output_format doit etre 'docx' ou 'txt'")

    content, out_filename, media_type = build_output_file(req.paragraphs, req.output_format)
    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
    )


# ---------------------------------------------------------------------------
# Routes admin
# ---------------------------------------------------------------------------

@app.post("/admin/languages", dependencies=[Depends(check_admin)])
def add_language(item: LanguageItem):
    langs = load_languages()
    if any(l["code"] == item.code for l in langs):
        raise HTTPException(status_code=400, detail="Cette langue est deja dans la liste")
    langs.append(item.dict())
    save_languages(langs)
    return {"status": "ajoutee", "languages": langs}


@app.delete("/admin/languages/{code}", dependencies=[Depends(check_admin)])
def remove_language(code: str):
    langs = [l for l in load_languages() if l["code"] != code]
    save_languages(langs)
    return {"status": "supprimee", "languages": langs}


@app.get("/admin/admins", dependencies=[Depends(check_admin)])
def get_admins():
    return load_admins()


@app.post("/admin/admins", dependencies=[Depends(check_admin)])
def add_admin(item: AdminEmailItem):
    data = load_admins()
    if item.email not in data["emails"]:
        data["emails"].append(item.email)
        save_admins(data)
    return data


@app.delete("/admin/admins/{email}", dependencies=[Depends(check_admin)])
def remove_admin(email: str):
    data = load_admins()
    data["emails"] = [e for e in data["emails"] if e != email]
    save_admins(data)
    return data
